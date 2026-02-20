"""
ClauseAI - Contract Analysis FastAPI Backend
Production-ready API with async background processing + frontend serving
"""

from fastapi import FastAPI, UploadFile, File, HTTPException, BackgroundTasks
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, HTMLResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel
from typing import Optional, List
from pathlib import Path
from enum import Enum
from datetime import datetime
import uuid
import json
import os
import sys
import glob
import traceback

# Ensure project root is on Python path for imports
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

# Create directories
Path("uploads").mkdir(exist_ok=True)
Path("results").mkdir(exist_ok=True)
Path("frontend").mkdir(exist_ok=True)


class StatusEnum(str, Enum):
    PENDING = "pending"
    PROCESSING = "processing"
    COMPLETED = "completed"
    FAILED = "failed"


class AnalysisResponse(BaseModel):
    """Response after upload"""
    task_id: str
    status: StatusEnum
    message: str


class AnalysisStatusResponse(BaseModel):
    """Response for status check"""
    task_id: str
    status: StatusEnum
    contract_type: Optional[str] = None
    industry: Optional[str] = None
    error: Optional[str] = None


class HistoryItem(BaseModel):
    task_id: str
    file_name: str
    status: str
    contract_type: Optional[str] = None
    industry: Optional[str] = None
    created_at: Optional[str] = None


app = FastAPI(title="ClauseAI", version="1.0.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


def run_analysis(task_id: str, file_path: str):
    """Background task - runs the graph"""
    metadata = {}
    try:
        # Update status to processing
        metadata_path = f"results/{task_id}.json"
        with open(metadata_path, "r") as f:
            metadata = json.load(f)
        metadata["status"] = "processing"
        with open(metadata_path, "w") as f:
            json.dump(metadata, f)
        
        # Import and run your graph
        from classifier import graph
        
        initial_state = {
            "file_path": file_path,
            "contract_text": "",
            "contract_type": None,
            "industry": None,
            "clauses": [],
            "clause_analysis": None,
            "review_plan": None,
            "sections": [],
            "modifications": [],
            "final_report": None
        }
        
        print(f"Starting contract analysis...")
        result = graph.invoke(initial_state)
        
        # Save final report as text file
        report_path = f"results/{task_id}_report.txt"
        with open(report_path, "w", encoding="utf-8") as f:
            f.write(result.get('final_report', 'Report generation failed'))
        
        # Update metadata with results
        metadata["status"] = "completed"
        metadata["contract_type"] = result.get('contract_type')
        metadata["industry"] = result.get('industry')
        with open(metadata_path, "w") as f:
            json.dump(metadata, f)
        
        print(f"✓ Analysis complete: {task_id}")
    
    except Exception as e:
        # Mark as failed
        print(f"Analysis failed: {task_id} - {e}")
        traceback.print_exc()
        metadata_path = f"results/{task_id}.json"
        try:
            with open(metadata_path, "r") as f:
                metadata = json.load(f)
        except Exception:
            metadata = {"task_id": task_id}
        metadata["status"] = "failed"
        metadata["error"] = str(e)
        with open(metadata_path, "w") as f:
            json.dump(metadata, f)


# ─── API ENDPOINTS ────────────────────────────────────────────

@app.post("/analyze", response_model=AnalysisResponse)
async def upload_contract(file: UploadFile = File(...), background_tasks: BackgroundTasks = BackgroundTasks()):
    """Upload contract and start analysis"""
    
    # Validate file type
    if not file.filename.endswith((".pdf", ".docx")):
        raise HTTPException(status_code=400, detail="Only PDF or DOCX files allowed")
    
    # Generate unique task ID
    task_id = str(uuid.uuid4())[:8]
    file_path = f"uploads/{task_id}_{file.filename}"
    
    # Save uploaded file
    with open(file_path, "wb") as f:
        f.write(await file.read())
    
    # Create metadata file
    metadata = {
        "task_id": task_id,
        "file_name": file.filename,
        "status": "pending",
        "contract_type": None,
        "industry": None,
        "error": None,
        "created_at": datetime.now().isoformat()
    }
    
    with open(f"results/{task_id}.json", "w") as f:
        json.dump(metadata, f)
    
    # Start background analysis
    background_tasks.add_task(run_analysis, task_id, file_path)
    
    return AnalysisResponse(
        task_id=task_id,
        status="pending",
        message="Analysis started in background"
    )


@app.get("/status/{task_id}")
async def get_status(task_id: str):
    """Check analysis status"""
    try:
        with open(f"results/{task_id}.json") as f:
            metadata = json.load(f)
        
        return AnalysisStatusResponse(**metadata)
    
    except FileNotFoundError:
        raise HTTPException(status_code=404, detail="Task not found")


@app.get("/result/{task_id}")
async def get_result(task_id: str):
    """Download final report"""
    report_path = Path(f"results/{task_id}_report.txt")
    if report_path.exists():
        return FileResponse(str(report_path), media_type="text/plain")
    raise HTTPException(status_code=404, detail="Report not ready or not found")


@app.get("/history")
async def get_history():
    """List all past analyses"""
    history = []
    for json_file in sorted(glob.glob("results/*.json"), key=os.path.getmtime, reverse=True):
        try:
            with open(json_file) as f:
                data = json.load(f)
            history.append({
                "task_id": data.get("task_id"),
                "file_name": data.get("file_name", "Unknown"),
                "status": data.get("status"),
                "contract_type": data.get("contract_type"),
                "industry": data.get("industry"),
                "created_at": data.get("created_at"),
            })
        except Exception:
            continue
    return history


@app.delete("/history/{task_id}")
async def delete_history_item(task_id: str):
    """Delete a specific analysis result"""
    json_path = Path(f"results/{task_id}.json")
    report_path = Path(f"results/{task_id}_report.txt")
    if not json_path.exists():
        raise HTTPException(status_code=404, detail="Task not found")
    
    json_path.unlink(missing_ok=True)
    report_path.unlink(missing_ok=True)
    
    # Also remove uploaded file
    for upload in Path("uploads").glob(f"{task_id}_*"):
        upload.unlink(missing_ok=True)
    
    return {"message": "Deleted successfully"}


@app.get("/health")
async def health_check():
    return {"status": "healthy", "version": "1.0.0"}


# ─── DOWNLOAD ENDPOINTS (PDF, DOCX, TXT) ─────────────────────

@app.get("/download/{task_id}")
async def download_report(task_id: str, format: str = "txt"):
    """Download report in various formats: txt, pdf, docx"""
    report_path = Path(f"results/{task_id}_report.txt")
    if not report_path.exists():
        raise HTTPException(status_code=404, detail="Report not found")

    report_text = report_path.read_text(encoding="utf-8")

    # Load metadata for title info
    meta = {}
    meta_path = Path(f"results/{task_id}.json")
    if meta_path.exists():
        with open(meta_path) as f:
            meta = json.load(f)

    if format == "txt":
        return FileResponse(
            str(report_path),
            media_type="text/plain",
            filename=f"ClauseAI_Report_{task_id}.txt"
        )

    elif format == "pdf":
        from fpdf import FPDF
        import tempfile

        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=20)
        pdf.add_page()

        # Title
        pdf.set_font("Helvetica", "B", 18)
        pdf.multi_cell(0, 12, "ClauseAI Contract Analysis Report", align="C")
        pdf.ln(4)

        # Meta line
        pdf.set_font("Helvetica", "", 10)
        pdf.set_text_color(100, 100, 100)
        contract_type = meta.get("contract_type", "Unknown")
        industry = meta.get("industry", "General")
        pdf.multi_cell(0, 6, f"Type: {contract_type}  |  Industry: {industry}  |  Task: {task_id}", align="C")
        pdf.ln(6)
        pdf.set_text_color(0, 0, 0)

        # Separator
        pdf.set_draw_color(200, 200, 200)
        pdf.line(15, pdf.get_y(), 195, pdf.get_y())
        pdf.ln(6)

        def safe_text(txt):
            """Sanitize text for fpdf - remove problematic chars"""
            return txt.encode('latin-1', errors='replace').decode('latin-1')

        # Body text
        for line in report_text.split("\n"):
            stripped = line.strip()
            if not stripped:
                pdf.ln(3)
                continue

            text = safe_text(stripped)

            # Headers
            if text.startswith("### "):
                pdf.set_font("Helvetica", "B", 12)
                pdf.set_x(pdf.l_margin)
                pdf.multi_cell(0, 6, text[4:])
                pdf.ln(2)
            elif text.startswith("## "):
                pdf.ln(3)
                pdf.set_font("Helvetica", "B", 14)
                pdf.set_x(pdf.l_margin)
                pdf.multi_cell(0, 7, text[3:])
                pdf.set_draw_color(200, 200, 200)
                pdf.line(15, pdf.get_y() + 1, 195, pdf.get_y() + 1)
                pdf.ln(4)
            elif text.startswith("# "):
                pdf.set_font("Helvetica", "B", 16)
                pdf.set_x(pdf.l_margin)
                pdf.multi_cell(0, 8, text[2:])
                pdf.ln(3)
            elif text.startswith("===="):
                pdf.set_draw_color(180, 180, 180)
                pdf.line(15, pdf.get_y(), 195, pdf.get_y())
                pdf.ln(4)
            elif text.startswith("- ") or text.startswith("* "):
                pdf.set_font("Helvetica", "", 10)
                pdf.set_x(pdf.l_margin + 8)
                w = pdf.epw - 8
                pdf.multi_cell(w, 5, "  " + text[2:])
            elif text.startswith("**") and text.endswith("**"):
                pdf.set_font("Helvetica", "B", 11)
                pdf.set_x(pdf.l_margin)
                pdf.multi_cell(0, 6, text.strip("*"))
            else:
                pdf.set_font("Helvetica", "", 10)
                clean = text.replace("**", "")
                pdf.set_x(pdf.l_margin)
                pdf.multi_cell(0, 5, clean)

        pdf_output = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
        pdf.output(pdf_output.name)
        pdf_output.close()

        return FileResponse(
            pdf_output.name,
            media_type="application/pdf",
            filename=f"ClauseAI_Report_{task_id}.pdf"
        )

    elif format == "docx":
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import tempfile

        doc = Document()

        # Style the default font
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(10)

        # Title
        title = doc.add_heading("ClauseAI Contract Analysis Report", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Meta
        contract_type = meta.get("contract_type", "Unknown")
        industry = meta.get("industry", "General")
        meta_para = doc.add_paragraph()
        meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = meta_para.add_run(f"Type: {contract_type}  |  Industry: {industry}  |  Task: {task_id}")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(130, 130, 130)

        doc.add_paragraph()  # spacing

        # Body
        for line in report_text.split("\n"):
            stripped = line.strip()
            if not stripped:
                continue
            if stripped.startswith("===="):
                doc.add_paragraph("_" * 60)
            elif stripped.startswith("### "):
                doc.add_heading(stripped[4:], level=3)
            elif stripped.startswith("## "):
                doc.add_heading(stripped[3:], level=2)
            elif stripped.startswith("# "):
                doc.add_heading(stripped[2:], level=1)
            elif stripped.startswith("- ") or stripped.startswith("* "):
                doc.add_paragraph(stripped[2:], style="List Bullet")
            else:
                p = doc.add_paragraph()
                # Parse **bold** sections
                import re as _re
                parts = _re.split(r'(\*\*[^*]+\*\*)', stripped)
                for part in parts:
                    if part.startswith("**") and part.endswith("**"):
                        run = p.add_run(part[2:-2])
                        run.bold = True
                    else:
                        p.add_run(part)

        docx_output = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
        doc.save(docx_output.name)
        docx_output.close()

        return FileResponse(
            docx_output.name,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename=f"ClauseAI_Report_{task_id}.docx"
        )

    else:
        raise HTTPException(status_code=400, detail="Format must be: txt, pdf, or docx")


# ─── SERVE FRONTEND ──────────────────────────────────────────

# Try React build first, then fallback to old frontend/
REACT_BUILD = Path("frontend-react/dist")
LEGACY_FRONTEND = Path("frontend")

if REACT_BUILD.exists():
    app.mount("/assets", StaticFiles(directory=str(REACT_BUILD / "assets")), name="assets")
    
    @app.get("/{full_path:path}")
    async def serve_react(full_path: str):
        """Serve React SPA - all routes go to index.html"""
        # Try to serve static file first
        file_path = REACT_BUILD / full_path
        if file_path.is_file():
            return FileResponse(str(file_path))
        # Fallback to index.html for SPA routing
        return FileResponse(str(REACT_BUILD / "index.html"))
else:
    app.mount("/static", StaticFiles(directory="frontend"), name="static")
    
    @app.get("/", response_class=HTMLResponse)
    async def serve_frontend():
        index_path = LEGACY_FRONTEND / "index.html"
        if index_path.exists():
            return index_path.read_text(encoding="utf-8")
        return HTMLResponse("<h1>Frontend not found. Run: cd frontend-react && npm run build</h1>", status_code=404)


if __name__ == "__main__":
    import uvicorn
    print("\n  ClauseAI running at: http://localhost:8000\n")
    uvicorn.run(app, host="0.0.0.0", port=8000)
