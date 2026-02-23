import os
import json
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from docx import Document
from .models import Analysis

async def export_report(a: Analysis, fmt: str) -> str:
    out_dir = os.path.join(os.getcwd(), "exports")
    os.makedirs(out_dir, exist_ok=True)
    base = os.path.splitext(a.filename)[0]
    if fmt == "txt":
        path = os.path.join(out_dir, base + ".txt")
        with open(path, "w", encoding="utf-8") as f:
            f.write(a.summary or "")
            if a.classification:
                f.write("\n\nClassification\n")
                f.write(a.classification)
            if a.risk_assessment:
                f.write("\n\nRisk\n")
                f.write(a.risk_assessment)
            if a.missing_clauses:
                f.write("\n\nMissing Clauses\n")
                f.write(a.missing_clauses)
            if a.experts_review:
                f.write("\n\nExperts Review\n")
                f.write(a.experts_review)
            if a.suggestions:
                f.write("\n\nSuggestions\n")
                f.write(a.suggestions)
        return path
    if fmt == "json":
        path = os.path.join(out_dir, base + ".json")
        with open(path, "w", encoding="utf-8") as f:
            f.write(a.json_result or "{}")
        return path
    if fmt == "docx":
        path = os.path.join(out_dir, base + ".docx")
        doc = Document()
        doc.add_heading("Analysis Report", level=1)
        doc.add_heading("Professional Summary", level=2)
        doc.add_paragraph(a.summary or "")
        if a.classification:
            doc.add_heading("Smart Classification", level=2)
            doc.add_paragraph(a.classification)
        if a.risk_assessment:
            doc.add_heading("Risk Assessment", level=2)
            doc.add_paragraph(a.risk_assessment)
        if a.missing_clauses:
            doc.add_heading("Missing Clauses", level=2)
            doc.add_paragraph(a.missing_clauses)
        if a.experts_review:
            doc.add_heading("Multi-Expert Review", level=2)
            doc.add_paragraph(a.experts_review)
        if a.suggestions:
            doc.add_heading("Smart Suggestions", level=2)
            doc.add_paragraph(a.suggestions)
        doc.save(path)
        return path
    path = os.path.join(out_dir, base + ".pdf")
    c = canvas.Canvas(path, pagesize=letter)
    width, height = letter
    y = height - 50
    def write(text, title=None):
        nonlocal y
        if title:
            c.setFont("Helvetica-Bold", 14)
            c.drawString(40, y, title)
            y -= 20
        c.setFont("Helvetica", 10)
        for line in (text or "").splitlines():
            c.drawString(40, y, line[:95])
            y -= 14
            if y < 60:
                c.showPage()
                y = height - 50
    write("Analysis Report")
    write(a.summary or "", "Professional Summary")
    if a.classification:
        write(a.classification, "Smart Classification")
    if a.risk_assessment:
        write(a.risk_assessment, "Risk Assessment")
    if a.missing_clauses:
        write(a.missing_clauses, "Missing Clauses")
    if a.experts_review:
        write(a.experts_review, "Multi-Expert Review")
    if a.suggestions:
        write(a.suggestions, "Smart Suggestions")
    c.save()
    return path
